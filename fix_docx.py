"""
修改 test_formatted.docx：
1. 全文删除序号后面的空格，如 "4. " -> "4."
2. 主事件流每条步骤前加角色标签 【管理员】/【系统】/【参会人员】 等
"""
from docx import Document
import re

doc = Document('test_formatted.docx')

# 每张表的主事件流角色分配
ROLE_MAP = {
    0: {  # UC-HYXT-HYGL-01 新建会议
        1: '【管理员】', 2: '【系统】', 3: '【管理员】', 4: '【系统】',
        5: '【管理员】', 6: '【系统】', 7: '【系统】', 8: '【管理员】',
        9: '【管理员】', 10: '【系统】'
    },
    1: {  # UC-HYXT-HYGL-02 编辑删除会议
        1: '【管理员】', 2: '【管理员】', 3: '【系统】', 4: '【管理员】',
        5: '【系统】', 6: '【系统】'
    },
    2: {  # UC-HYXT-HYGL-03 管理会务材料
        1: '【管理员】', 2: '【系统】', 3: '【管理员】', 4: '【系统】',
        5: '【管理员】', 6: '【系统】', 7: '【管理员】', 8: '【系统】'
    },
    3: {  # UC-HYXT-HYJL-01 实时语音转写
        1: '【管理员】', 2: '【系统】', 3: '【系统】', 4: '【系统】',
        5: '【系统】', 6: '【系统】', 7: '【系统】', 8: '【系统】',
        9: '【管理员】', 10: '【系统】'
    },
    4: {  # UC-HYXT-HYJL-02 优化校验转写
        1: '【系统】', 2: '【系统】', 3: '【系统】', 4: '【管理员】',
        5: '【系统】', 6: '【管理员】', 7: '【系统】', 8: '【管理员】',
        9: '【系统】'
    },
    5: {  # UC-HYXT-HYJL-03 导出会议记录
        1: '【管理员】', 2: '【系统】', 3: '【管理员】', 4: '【管理员】',
        5: '【系统】', 6: '【系统】'
    },
    6: {  # UC-HYXT-HYJL-04 AI智能问答（主要角色：管理员/参会人员）
        1: '【管理员/参会人员】', 2: '【系统】', 3: '【管理员/参会人员】',
        4: '【管理员/参会人员】', 5: '【系统】', 6: '【系统】',
        7: '【系统】', 8: '【系统】'
    },
    7: {  # UC-HYXT-YDSC-01 生成会议要点
        1: '【管理员】', 2: '【系统】', 3: '【管理员】', 4: '【系统】',
        5: '【系统】', 6: '【系统】', 7: '【系统】', 8: '【系统】',
        9: '【管理员】'
    },
    8: {  # UC-HYXT-YDSC-02 标注会议记录
        1: '【管理员】', 2: '【管理员】', 3: '【系统】', 4: '【管理员】',
        5: '【系统】', 6: '【系统】', 7: '【管理员】'
    },
    9: {  # UC-HYXT-JYSC-01 生成会议纪要
        1: '【管理员】', 2: '【管理员】', 3: '【系统】', 4: '【系统】',
        5: '【系统】', 6: '【系统】', 7: '【系统】', 8: '【系统】'
    },
    10: {  # UC-HYXT-JYSC-02 编辑播报纪要
        1: '【管理员】', 2: '【管理员】', 3: '【管理员】', 4: '【系统】',
        5: '【管理员】', 6: '【管理员】', 7: '【系统】', 8: '【管理员】',
        9: '【系统】'
    },
    11: {  # UC-HYXT-DZSQ-01 参会人员签到
        1: '【参会人员】', 2: '【系统】', 3: '【系统】', 4: '【参会人员】',
        5: '【系统】', 6: '【参会人员】', 7: '【参会人员】', 8: '【系统】',
        9: '【系统】', 10: '【系统】'
    },
    12: {  # UC-HYXT-DZSQ-02 审签会议纪要
        1: '【管理员】', 2: '【系统】', 3: '【管理员】', 4: '【系统】',
        5: '【签署人】', 6: '【签署人】', 7: '【签署人】', 8: '【系统】',
        9: '【系统】'
    },
    13: {  # UC-HYXT-DBSX-01 分发待办事项
        1: '【管理员】', 2: '【系统】', 3: '【管理员】', 4: '【管理员】',
        5: '【管理员】', 6: '【系统】', 7: '【系统】', 8: '【责任人】',
        9: '【系统】'
    },
    14: {  # UC-HYXT-JDGZ-01 跟踪待办进度（主要角色：管理员/责任人）
        1: '【管理员/责任人】', 2: '【系统】', 3: '【系统】', 4: '【系统】',
        5: '【系统】', 6: '【责任人】', 7: '【系统】', 8: '【责任人】',
        9: '【系统】', 10: '【系统】'
    },
    15: {  # UC-HYXT-JDGZ-02 归档会议数据
        1: '【管理员】', 2: '【管理员】', 3: '【系统】', 4: '【系统】',
        5: '【系统】', 6: '【系统】', 7: '【用户】', 8: '【系统】',
        9: '【用户】'
    },
}


def remove_number_space_in_run(text: str) -> str:
    """将 'N. ' 模式（序号+句点+空格）中的空格删除。"""
    return re.sub(r'(\d+)\. ', r'\1.', text)


def add_role_tags_to_text(text: str, role_map: dict) -> str:
    """
    先删除序号后空格，再向每条编号步骤前插入角色标签。
    输入格式：'1. 管理员…。2. 系统…。'
    输出格式：'1.【管理员】管理员…。2.【系统】系统…。'
    """
    # Step1: 去除序号后空格
    text = re.sub(r'(\d+)\. ', r'\1.', text)

    # Step2: 定位第一个步骤编号的起始位置
    first = re.search(r'\d+\.', text)
    if not first:
        return text

    prefix = text[:first.start()]
    rest = text[first.start():]

    # Step3: 按 "N." 分割步骤
    steps = re.findall(r'(\d+)\.(.*?)(?=\d+\.|$)', rest, re.DOTALL)
    if not steps:
        return text

    result = [prefix]
    for num_str, content in steps:
        num = int(num_str)
        role = role_map.get(num, '')
        result.append(f'{num_str}.{role}{content}')
    return ''.join(result)


def process_runs_remove_space(para):
    """对段落的每个 run 执行序号空格删除。"""
    for run in para.runs:
        if run.text:
            new_text = remove_number_space_in_run(run.text)
            if new_text != run.text:
                run.text = new_text


def process_main_flow_para(para, role_map: dict):
    """
    处理主事件流单元格段落：合并所有 run 文本，添加角色标签后写回。
    仅在第一个 run 中保留全部新文本，其余 run 清空。
    """
    full_text = ''.join(run.text for run in para.runs)
    if not full_text.strip():
        return
    new_text = add_role_tags_to_text(full_text, role_map)
    if new_text == full_text:
        return
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)


# ── 处理正文段落（非表格） ──────────────────────────────────────────────────
for para in doc.paragraphs:
    process_runs_remove_space(para)

# ── 处理表格单元格 ─────────────────────────────────────────────────────────
for t_idx, table in enumerate(doc.tables):
    for row in table.rows:
        is_main_flow_row = row.cells[0].text.strip() == '主事件流'
        for c_idx, cell in enumerate(row.cells):
            is_main_flow_cell = (
                is_main_flow_row and c_idx == 1 and t_idx in ROLE_MAP
            )
            for para in cell.paragraphs:
                if is_main_flow_cell:
                    process_main_flow_para(para, ROLE_MAP[t_idx])
                else:
                    process_runs_remove_space(para)

doc.save('test_formatted.docx')
print('完成！文件已保存：test_formatted.docx')
