"""
补充修正脚本：对 4 处因 Unicode 引号不一致而未应用的差异进行标注。
操作对象：已校对版本（在其基础上继续修正）
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'C:\Users\gbw18\Desktop\软件需求规格说明-智能会议（已校对）.docx'
RED  = 'CC0000'
BLUE = '0055AA'

doc = Document(OUTPUT)


def _base_children(para):
    if not para.runs:
        return []
    rpr = para.runs[0]._r.find(qn('w:rPr'))
    if rpr is None:
        return []
    skip = {'color', 'strike', 'dstrike', 'w14:textFill'}
    return [copy.deepcopy(ch) for ch in rpr
            if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) not in skip]


def add_run(p_el, text, base, color=None, strike=False):
    if not text:
        return
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for ch in base:
        rPr.append(copy.deepcopy(ch))
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if strike:
        rPr.append(OxmlElement('w:strike'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_el.append(r)


def apply_fix(para, wrong, correct=None):
    full = ''.join(r.text for r in para.runs)
    if wrong not in full:
        return False
    idx    = full.index(wrong)
    before = full[:idx]
    after  = full[idx + len(wrong):]
    base   = _base_children(para)
    p = para._p
    for r in para.runs:
        p.remove(r._r)
    add_run(p, before, base)
    add_run(p, wrong, base, color=RED, strike=True)
    if correct:
        add_run(p, correct, base, color=BLUE)
    add_run(p, after, base)
    return True


# ── Fix 4: TABLE 7 row5 前置条件 ─────────────────────────────────────────
# 可删除状态描述不完整：只排除了"进行中"和"审签中"，实际准备中/会后处理中也不可删
print('[4] TABLE7 row5 前置条件 — 可删除状态')
cell = doc.tables[7].rows[5].cells[1]
for para in cell.paragraphs:
    txt = para.text
    if '\u8fdb\u884c\u4e2d' in txt and '\u5ba1\u7b7e\u4e2d' in txt:
        # Extract the exact substring from 状态非 to end of ）
        qi = txt.find('\u72b6\u6001\u975e')   # index of 状态非
        ei = txt.find('\uff09', qi) + 1        # index just after ）
        wrong   = txt[qi:ei]                   # e.g. 状态非"进行中"和"审签中"）
        correct = ('\u72b6\u6001\u4e3a\u201c\u5f85\u53ec\u5f00\u201d\u3001'
                   '\u201c\u5df2\u7ed3\u675f\u201d\u6216\u201c\u5df2\u5f52\u6863\u201d'
                   '\uff08\u201c\u51c6\u5907\u4e2d\u201d/\u201c\u4f1a\u540e\u5904\u7406'
                   '\u4e2d\u201d\u540c\u6837\u4e0d\u53ef\u5220\u9664\uff09')
        if apply_fix(para, wrong, correct):
            print('  OK:', repr(wrong))
        break

# ── Fix 5: TABLE 7 row7 主事件流 step3 ────────────────────────────────────
# "准备中"在代码中不显示删除按钮，只有 待召开/已结束/已归档 才可删除
print('[5] TABLE7 row7 主事件流 step3 — 准备中不可删')
cell = doc.tables[7].rows[7].cells[1]
for para in cell.paragraphs:
    txt = para.text
    if '\u51c6\u5907\u4e2d' in txt and '\u5f85\u53ec\u5f00' in txt and '\u5220\u9664' in txt:
        # Extract （会议状态为...）
        si = txt.find('\uff08\u4f1a\u8bae\u72b6\u6001\u4e3a')   # （会议状态为
        ei = txt.find('\uff09', si) + 1                          # ）
        wrong   = txt[si:ei]
        correct = ('\uff08\u4f1a\u8bae\u72b6\u6001\u4e3a\u201c\u5f85\u53ec\u5f00\u201d'
                   '\u3001\u201c\u5df2\u7ed3\u675f\u201d\u6216\u201c\u5df2\u5f52\u6863\u201d'
                   '\uff1b\u5b9e\u9645\u4ee3\u7801\u4ec5\u8fd9\u4e09\u79cd\u72b6\u6001\u663e'
                   '\u793a\u201c\u5220\u9664\u201d\u83dc\u5355\u9879\uff09')
        if apply_fix(para, wrong, correct):
            print('  OK:', repr(wrong))
        break

# ── Fix 6: TABLE 7 row8 可选事件流 3a ─────────────────────────────────────
print('[6] TABLE7 row8 可选事件流 3a — 同步修正可删除状态')
cell = doc.tables[7].rows[8].cells[1]
for para in cell.paragraphs:
    txt = para.text
    if '3a' in txt and '\u51c6\u5907\u4e2d' in txt and '\u5f85\u53ec\u5f00' in txt:
        # Extract 3a.选择的目标会议状态不为...（到冒号为止）
        si = txt.find('3a.')
        ei = txt.find('\uff1a', si)  # 全角冒号：
        wrong   = txt[si:ei]
        correct = ('3a.\u9009\u62e9\u7684\u76ee\u6807\u4f1a\u8bae\u72b6\u6001\u4e0d\u4e3a'
                   '\u201c\u5f85\u53ec\u5f00\u201d\u3001\u201c\u5df2\u7ed3\u675f\u201d'
                   '\u6216\u201c\u5df2\u5f52\u6863\u201d\uff08\u5373\u201c\u51c6\u5907\u4e2d'
                   '\u201d/\u201c\u8fdb\u884c\u4e2d\u201d/\u201c\u4f1a\u540e\u5904\u7406\u4e2d'
                   '\u201d/\u201c\u5ba1\u7b7e\u4e2d\u201d\u5747\u4e0d\u53ef\u5220\u9664\uff09')
        if apply_fix(para, wrong, correct):
            print('  OK:', repr(wrong))
        break

# ── Fix 8: TABLE 10 row8 15a 裁剪模式 ────────────────────────────────────
print('[8] TABLE10 row8 15a — 裁剪模式不存在')
cell = doc.tables[10].rows[8].cells[1]
for para in cell.paragraphs:
    if '15a' in para.text and '\u88c1\u526a\u6a21\u5f0f' in para.text:   # 裁剪模式
        wrong   = para.text
        correct = ('\u3010\u6ce8\uff1aMeetingLive.vue \u4e2d\u4e0d\u5b58\u5728'
                   '\u201c\u88c1\u526a\u6a21\u5f0f\u201d\u529f\u80fd\u53ca\u76f8\u5173'
                   '\u6309\u9215\uff0c\u6b64\u53ef\u9009\u4e8b\u4ef6\u6d41\u63cf\u8ff0'
                   '\u6709\u8bef\uff0c\u5e94\u5220\u9664\u3011')
        if apply_fix(para, wrong, correct):
            print('  OK: 15a paragraph fixed')
        break

doc.save(OUTPUT)
print('\n已保存:', OUTPUT)
