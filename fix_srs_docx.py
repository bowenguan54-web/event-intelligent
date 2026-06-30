"""
校正 软件需求规格说明-智能会议.docx
--------------------------------------
根据项目实际代码，将文档中与实现不符的内容用颜色标注：
  错误/废弃文本 → 红色 + 删除线 (CC0000)
  修正文本      → 蓝色 (0055AA)
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r'C:\Users\gbw18\Desktop\软件需求规格说明-智能会议.docx'
OUTPUT = r'C:\Users\gbw18\Desktop\软件需求规格说明-智能会议（已校对）.docx'

RED  = 'CC0000'   # 错误文本颜色
BLUE = '0055AA'   # 修正文本颜色

doc = Document(INPUT)


# ──────────────────────────────────────────────
# 核心工具函数
# ──────────────────────────────────────────────

def _base_rpr_children(para):
    """从段落第一个 run 提取基础格式（字号等），剥离颜色/删除线。"""
    if not para.runs:
        return []
    rpr = para.runs[0]._r.find(qn('w:rPr'))
    if rpr is None:
        return []
    skip = {'color', 'strike', 'dstrike', 'w14:textFill'}
    return [copy.deepcopy(ch) for ch in rpr
            if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) not in skip]


def _add_run(p_element, text, base_children, color_hex=None, strike=False):
    """在 <w:p> 末尾追加一个有格式的 run。"""
    if not text:
        return
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for ch in base_children:
        rPr.append(copy.deepcopy(ch))
    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        rPr.append(c)
    if strike:
        rPr.append(OxmlElement('w:strike'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_element.append(r)


def fix_para(para, wrong, correct=None):
    """
    在 para 中找到 wrong，标红删除线；
    若提供 correct，紧接其后以蓝色插入。
    返回 True 表示成功匹配并修改。
    """
    full = ''.join(r.text for r in para.runs)
    if wrong not in full:
        return False

    idx    = full.index(wrong)
    before = full[:idx]
    after  = full[idx + len(wrong):]
    base   = _base_rpr_children(para)

    p = para._p
    for r in para.runs:
        p.remove(r._r)

    _add_run(p, before, base)
    _add_run(p, wrong,   base, color_hex=RED,  strike=True)
    if correct:
        _add_run(p, correct, base, color_hex=BLUE)
    _add_run(p, after, base)
    return True


def fix_in_cell(tbl_idx, row_idx, col_idx, wrong, correct=None):
    """在指定表格单元格中找到 wrong 并修正；遍历格内所有段落。"""
    cell = doc.tables[tbl_idx].rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        if fix_para(para, wrong, correct):
            return True
    print(f'  [未找到] table={tbl_idx} row={row_idx} col={col_idx} text="{wrong[:30]}..."')
    return False


def fix_in_doc_para(para_idx, wrong, correct=None):
    """在文档段落列表中找到 wrong 并修正。"""
    if not fix_para(doc.paragraphs[para_idx], wrong, correct):
        print(f'  [未找到] para={para_idx} text="{wrong}"')


# ──────────────────────────────────────────────
# 开始修正
# ──────────────────────────────────────────────

print('开始校对文档...\n')

# ── 1. 文档标题段落：YDSD-02 → YDSC-02（排版错误） ──────────────────────
print('[ 1] 标注会议记录 UC 编号 YDSD→YDSC（正文标题段落）')
for i, p in enumerate(doc.paragraphs):
    if 'YDSD-02' in p.text:
        fix_in_doc_para(i, 'YDSD-02', 'YDSC-02')
        print(f'     → 已修正段落 {i}')
        break

# ── 2. TABLE 5 (UC-HYXT-HYGL-01 新建会议) ─────────────────────────────────
# 主事件流 step4：显示步骤条有 ⑤智能议程，但向导实际只有4步
print('\n[ 2] TABLE 5 新建会议 主事件流 — 删除不存在的"⑤智能议程"步骤')
fix_in_cell(5, 7, 1,
    '、⑤智能议程',
    '（注：实际向导仅4步，无"智能议程"第5步）')

# ── 3. TABLE 6 (UC-HYXT-HYGL-02 编辑会议) ────────────────────────────────
# 主事件流 step4："步骤条与新建会议一致（5步）" → 实为4步
print('\n[ 3] TABLE 6 编辑会议 主事件流 — 步骤数 5→4')
fix_in_cell(6, 7, 1,
    '步骤条与新建会议一致（5步）',
    '步骤条与新建会议一致（4步）')

# ── 4. TABLE 7 (UC-HYXT-HYGL-03 删除会议) ────────────────────────────────
# 前置条件：可删除状态描述不完整（preparing/processing也不可删）
print('\n[ 4] TABLE 7 删除会议 前置条件 — 可删除状态说明')
fix_in_cell(7, 5, 1,
    '状态非"进行中"和"审签中"',
    '状态为"待召开"、"已结束"或"已归档"（"准备中"/"会后处理中"同样不可删除）')

# 主事件流 step3：准备中的会议实际不显示删除按钮
print('\n[ 5] TABLE 7 删除会议 主事件流 — 准备中状态不可删')
fix_in_cell(7, 7, 1,
    '（会议状态为"待召开"或"准备中"）',
    '（会议状态为"待召开"、"已结束"或"已归档"；实际代码仅这三种状态显示"删除"菜单项）')

# 可选事件流 3a：与主事件流同步修正
print('\n[ 6] TABLE 7 删除会议 可选事件流 3a — 同步修正可删除状态')
fix_in_cell(7, 8, 1,
    '3a.选择的目标会议状态不为"待召开"或"准备中"',
    '3a.选择的目标会议状态不为"待召开"、"已结束"或"已归档"（即"准备中"/"进行中"/"会后处理中"/"审签中"）')

# ── 5. TABLE 8 (UC-HYXT-HYGL-04 管理会务材料) ────────────────────────────
# 主事件流：步骤数同样错写为5步
print('\n[ 7] TABLE 8 管理会务材料 主事件流 — 步骤数 5→4')
fix_in_cell(8, 7, 1,
    '步骤条与新建会议一致（5步）',
    '步骤条与新建会议一致（4步）')

# ── 6. TABLE 10 (UC-HYXT-HYJL-02 优化校验转写) ───────────────────────────
# 可选事件流 15a："裁剪模式" — 该功能在代码中完全不存在
print('\n[ 8] TABLE 10 优化校验转写 可选事件流 15a — 裁剪模式不存在')
fix_in_cell(10, 8, 1,
    '15a.管理员点击"进入裁剪模式"按钮：每段记录右侧出现"×"标记按钮，管理员点击"×"标记需删除的段落，点击"确认裁剪并归档"后系统批量删除已标记段落。',
    '【注：MeetingLive.vue 中不存在"裁剪模式"功能及相关按钮，此可选事件流描述有误，应删除】')

# ── 7. TABLE 12 (UC-HYXT-HYJL-04 AI智能问答) ─────────────────────────────
# 前置条件 1："用户已登录系统" — 终端参会人员通过IP识别，无账号登录
print('\n[ 9] TABLE 12 AI智能问答 前置条件 — 终端参会人员不通过账号登录')
fix_in_cell(12, 5, 1,
    '用户已登录系统',
    '参会人员通过会议终端IP自动识别身份（终端用户无需账号密码登录）')

# ── 8. TABLE 17 (UC-HYXT-DZSQ-01 参会人员签到) ───────────────────────────
# 描述行：会议码 → IP识别
print('\n[10] TABLE 17 参会人员签到 描述行 — 会议码→IP识别')
fix_in_cell(17, 4, 1,
    '通过会议码加入会议',
    '通过会议终端IP自动识别加入会议')

# 主事件流 step1：输入6位会议码 → IP自动识别（IP识别不需要手动输入）
print('\n[11] TABLE 17 参会人员签到 主事件流 step1 — 会议码输入→IP自动识别')
fix_in_cell(17, 7, 1,
    '输入6位会议码，点击"加入会议"',
    '进入会议终端（系统通过终端机IP自动识别对应会议及参会人员信息，无需手动输入会议码）')

# 主事件流 step2：会议码验证 → IP自动加载
print('\n[12] TABLE 17 参会人员签到 主事件流 step2 — 会议码验证→IP加载')
fix_in_cell(17, 7, 1,
    '系统验证会议码有效且会议处于"准备中"状态',
    '系统通过IP加载对应会议，展示身份确认面板（含姓名、单位、"已识别"标识）及手写签名区域')

# 可选事件流 2a：会议码无效（IP机制下不适用）
print('\n[13] TABLE 17 参会人员签到 可选事件流 2a — 会议码无效不适用')
fix_in_cell(17, 8, 1,
    '2a. 会议码无效：系统提示"会议码错误，请重新输入"。',
    '【注：系统通过IP识别，不使用会议码，此条不适用；实际错误为IP未匹配到对应座位/会议】')

# 可选事件流 2b：会议不在准备中（IP机制下错误类型不同）
print('\n[14] TABLE 17 参会人员签到 可选事件流 2b — 错误描述不适用')
fix_in_cell(17, 8, 1,
    '2b. 会议不在"准备中"状态：系统提示"当前不在签到时段"。',
    '【注：此条应改为"2b. 终端IP未能匹配有效座位：系统提示相应错误信息"】')

# ── 9. TABLE 20 (UC-HYXT-JDGZ-01 跟踪待办进度) ───────────────────────────
# 主事件流 steps 6-9："更新状态"按钮及弹窗在MeetingTrack.vue中均不存在
print('\n[15] TABLE 20 跟踪待办进度 主事件流 step6 — 无"更新状态"按钮')
fix_in_cell(20, 7, 1,
    '点击"更新状态"',
    '（注：MeetingTrack.vue 中不存在"更新状态"按钮；以下 step7~9 描述的交互均未实现，待办状态由外部系统接口 IF_LCGL_HYXT_WCQK 自动同步）')

print('\n[16] TABLE 20 跟踪待办进度 主事件流 step7 — 状态选择面板不存在')
fix_in_cell(20, 7, 1,
    '系统弹出状态选择面板（待处理/进行中/已完成）',
    '【此步骤不存在于实际系统中，应删除】')

print('\n[17] TABLE 20 跟踪待办进度 主事件流 step8 — 手动选择状态不存在')
fix_in_cell(20, 7, 1,
    '责任人选择新状态，点击"确认"',
    '【此步骤不存在于实际系统中，应删除】')

print('\n[18] TABLE 20 跟踪待办进度 主事件流 step9 — 变更日志自动记录而非手动触发')
fix_in_cell(20, 7, 1,
    '系统记录变更日志，若标记完成则记录完成时间',
    '【实际由外部接口自动触发状态同步，非手动操作产生；描述有误】')

# ──────────────────────────────────────────────
# 保存
# ──────────────────────────────────────────────
doc.save(OUTPUT)
print(f'\n✓ 已保存至：{OUTPUT}')
